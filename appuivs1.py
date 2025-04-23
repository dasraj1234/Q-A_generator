import streamlit as st
from PIL import Image
import pytesseract
from openai import OpenAI
from docx import Document
from io import BytesIO
import os
from dotenv import load_dotenv
import shutil
#tesseract_path = shutil.which("tesseract")
#if tesseract_path is None:
    #st.error("Tesseract is not installed or not in PATH.")
#else:
    #st.success(f"Tesseract found at: {tesseract_path}")


# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Function to perform OCR on an image and return extracted text
def ocr_extract_text(image_path):
    """
    This function takes an image file, processes it with Tesseract OCR,
    and returns the extracted text.
    """
    image = Image.open(image_path)
    pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"  # Default on Render Debian-based images
    text = pytesseract.image_to_string(image)
    return text

# Function to call GPT-4 for generating questions
def ask_gpt(prompt, max_tokens=1000):
    """
    This function interacts with OpenAI's GPT-4 model to generate responses
    based on the prompt provided. It returns the generated content.
    """
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are an expert question generator for academic purposes."},
                  {"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=max_tokens
    )
    return response.choices[0].message.content.strip()

# Function to identify topics and concepts from text
def identify_topics_and_concepts(text):
    """
    This function extracts the main topics and concepts from the given text.
    """
    prompt = f"Identify the main topics and concepts from the following text:\n\n{text}"
    return ask_gpt(prompt, max_tokens=500)

# Function to generate academic questions based on the text
def generate_questions(text):
    """
    This function generates various types of questions (MCQs, True/False, Fill-in-the-blanks, etc.)
    based on the provided text using GPT-4.
    """
    prompt = f"Generate a set of academic questions (MCQs, True/False, Fill-in-the-blanks, Descriptive) based on the following text:\n\n{text}"
    return ask_gpt(prompt, max_tokens=1000)

# Function to generate answers for the questions
def generate_answers(questions):
    """
    This function generates detailed answers and explanations for the given set of questions.
    """
    prompt = f"Provide detailed answers and explanations for the following questions:\n\n{questions}"
    return ask_gpt(prompt, max_tokens=1500)

# Function to create a well-structured Word document (.docx) with questions and answers
def create_word_document(questions, answers):
    """
    This function takes the generated questions and answers, structures them in a readable format,
    and saves them to a Word document.
    """
    doc = Document()
    doc.add_heading('Questions and Answers', 0)

    # Combine questions and answers in bullet point format
    content = f"**Questions:**\n{questions}\n\n**Answers:**\n{answers}"

    # Adding the questions and answers as bullet points
    doc.add_paragraph(content, style='ListBullet')
    
    return doc

# Main Streamlit app
def main():
    """
    This is the main function that runs the Streamlit UI.
    It handles image uploads, processes OCR, generates questions and answers, 
    and allows the user to download the result as a Word document.
    """
    st.title("AI-Driven Question and Answer Generator")

    # Step 1: Upload image for OCR
    uploaded_image = st.file_uploader("Upload Image for OCR", type=["jpg", "png", "jpeg"])
    
    if uploaded_image is not None:
        # Step 2: Process the image to extract text using OCR
        extracted_text = ocr_extract_text(uploaded_image)
        
        # Step 3: Identify topics and concepts from the extracted text
        topics_and_concepts = identify_topics_and_concepts(extracted_text)
        
        # Step 4: Generate questions based on the extracted text
        questions = generate_questions(extracted_text)
        
        # Step 5: Generate answers for the generated questions
        answers = generate_answers(questions)

        # Combine the questions and answers into one structured format for display
        structured_data = f"**Questions:**\n{questions}\n\n**Answers:**\n{answers}"

        # Step 6: Show all generated questions and answers in a single text area
        st.write("Generated Questions and Answers:")
        st.text_area("Questions and Answers", structured_data, height=400)

        # Step 7: Create a button for generating a Word document
        if st.button("Generate and Download DOCX"):
            # Create the document with questions and answers
            doc = create_word_document(questions, answers)

            # Save the document to a BytesIO stream to facilitate download
            byte_stream = BytesIO()
            doc.save(byte_stream)
            byte_stream.seek(0)

            # Offer the document for download
            st.download_button(
                label="Download DOCX",
                data=byte_stream,
                file_name="questions_and_answers.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Run the Streamlit app
if __name__ == "__main__":
    main()
